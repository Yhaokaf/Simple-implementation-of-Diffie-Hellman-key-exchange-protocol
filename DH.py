import sympy
import streamlit as st
from docx import Document
import io

# 生成一个100-255之间的素数
def get_prime_in_range(start=100, end=255):
    for num in range(start, end + 1):
        if sympy.isprime(num):
            return num
    return None

# 获取原根
def find_primitive_root(p):
    for g in range(2, p):
        if all(pow(g, (p-1)//f, p) != 1 for f in sympy.factorint(p-1).keys()):
            return g
    return None

# D-H密钥交换过程
def diffie_hellman_exchange(p, g, a, b):
    # 计算A和B的公钥
    A = pow(g, a, p)
    B = pow(g, b, p)

    # 计算共享密钥
    s_A = pow(B, a, p)
    s_B = pow(A, b, p)

    return A, B, s_A, s_B

# 记录计算过程到Word文档
def record_process_to_word(p, g, a, b, A, B, s_A, s_B):
    doc = Document()
    doc.add_heading('Diffie-Hellman Key Exchange', 0)

    doc.add_heading('Step 1: Choose Prime and Primitive Root', level=1)
    doc.add_paragraph(f"Prime p = {p}")
    doc.add_paragraph(f"Primitive Root g = {g}")

    doc.add_heading('Step 2: Choose Private Keys', level=1)
    doc.add_paragraph(f"Private key of A (a) = {a}")
    doc.add_paragraph(f"Private key of B (b) = {b}")

    doc.add_heading('Step 3: Calculate Public Keys', level=1)
    doc.add_paragraph(f"Public key of A (A) = {A}")
    doc.add_paragraph(f"Public key of B (B) = {B}")

    doc.add_heading('Step 4: Calculate Shared Key', level=1)
    doc.add_paragraph(f"Shared key calculated by A (s_A) = {s_A}")
    doc.add_paragraph(f"Shared key calculated by B (s_B) = {s_B}")

    doc.add_paragraph("Since s_A == s_B, the key exchange is successful.")

    # 保存到一个内存中的字节流
    byte_io = io.BytesIO()
    doc.save(byte_io)
    byte_io.seek(0)
    return byte_io

# Streamlit Web界面
def main():
    st.title('Diffie-Hellman Key Exchange Demo')

    # Step 1: 选择素数p和原根g
    p = get_prime_in_range()
    g = find_primitive_root(p)

    # Step 2: 用户输入私钥
    a = st.number_input("Private key of A (a):", min_value=1, max_value=p-1, value=sympy.randprime(1, p-1))
    b = st.number_input("Private key of B (b):", min_value=1, max_value=p-1, value=sympy.randprime(1, p-1))

    st.write(f"Prime p = {p}")
    st.write(f"Primitive Root g = {g}")
    st.write(f"Private key of A (a) = {a}")
    st.write(f"Private key of B (b) = {b}")

    if st.button('Calculate Exchange'):
        # Step 3: 执行D-H密钥交换
        A, B, s_A, s_B = diffie_hellman_exchange(p, g, a, b)

        # 展示结果
        st.write(f"Public key of A (A) = {A}")
        st.write(f"Public key of B (B) = {B}")
        st.write(f"Shared key calculated by A (s_A) = {s_A}")
        st.write(f"Shared key calculated by B (s_B) = {s_B}")
        if s_A == s_B:
            st.write("The key exchange is successful! The shared key is the same.")
        else:
            st.write("The key exchange failed. The shared keys are different.")

        # Step 4: 生成Word文档
        byte_io = record_process_to_word(p, g, a, b, A, B, s_A, s_B)

        # 提供下载链接
        st.download_button(
            label="Download Calculation Process as Word Document",
            data=byte_io,
            file_name="dh_key_exchange.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

if __name__ == "__main__":
    main()
