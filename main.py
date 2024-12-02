import sympy
import streamlit as st
from docx import Document
import io

# 生成一个100-255之间的素数
def get_prime_in_range(start=100, end=255):
    for num in range(start, end + 1):
        if sympy.isprime(num):
            return num
    return None

# 获取原根
def find_primitive_root(p):
    for g in range(2, p):
        if all(pow(g, (p-1)//f, p) != 1 for f in sympy.factorint(p-1).keys()):
            return g
    return None

# D-H密钥交换过程
def diffie_hellman_exchange(p, g, a, b):
    # 计算A和B的公钥
    A = pow(g, a, p)
    B = pow(g, b, p)

    # 计算共享密钥
    s_A = pow(B, a, p)
    s_B = pow(A, b, p)

    return A, B, s_A, s_B



from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import io



def record_process_to_word(p, g, a, b, A, B, s_A, s_B):
    doc = Document()


    # 标题
    title = doc.add_heading(level=0)
    title_run = title.add_run('Diffie-Hellman Key Exchange Process')
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 102, 204)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 简介
    intro_paragraph = doc.add_paragraph(
        'This document describes the process of the Diffie-Hellman key exchange protocol with detailed calculations and results.',
        style='Normal')
    intro_run = intro_paragraph.runs[0]
    intro_run.font.size = Pt(12)
    intro_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 添加创作者信息
    doc.add_paragraph("Creator: 20221587 史亚涛, 20221485 李昊轩", style='Normal')
    doc.add_paragraph("Instructor: 陈欣", style='Normal')

    # 步骤 1: 选择素数和原根
    doc.add_heading('Step 1: Choose Prime and Primitive Root', level=1)
    doc.add_paragraph(f"Prime p = {p}", style='BodyText')
    doc.add_paragraph(f"Primitive Root g = {g}", style='BodyText')

    # 步骤 2: 选择私钥
    doc.add_heading('Step 2: Choose Private Keys', level=1)
    doc.add_paragraph(f"Private key of A (a) = {a}", style='BodyText')
    doc.add_paragraph(f"Private key of B (b) = {b}", style='BodyText')

    # 步骤 3: 计算公钥
    doc.add_heading('Step 3: Calculate Public Keys', level=1)
    doc.add_paragraph(f"Public key of A (A) is calculated as:", style='BodyText')
    # 插入公式
    formula_A = f"A = g^a mod p = {g}^{a} mod {p} = {A}"
    formula_A_paragraph = doc.add_paragraph(formula_A, style='Normal')
    formula_A_run = formula_A_paragraph.runs[0]
    formula_A_run.bold = True

    doc.add_paragraph(f"Public key of B (B) is calculated as:", style='BodyText')
    formula_B = f"B = g^b mod p = {g}^{b} mod {p} = {B}"
    formula_B_paragraph = doc.add_paragraph(formula_B, style='Normal')
    formula_B_run = formula_B_paragraph.runs[0]
    formula_B_run.bold = True

    # 步骤 4: 计算共享密钥
    doc.add_heading('Step 4: Calculate Shared Key', level=1)
    doc.add_paragraph(f"Shared key calculated by A (s_A) is:", style='BodyText')
    # 插入公式
    formula_s_A = f"s_A = B^a mod p = {B}^{a} mod {p} = {s_A}"
    formula_s_A_paragraph = doc.add_paragraph(formula_s_A, style='Normal')
    formula_s_A_run = formula_s_A_paragraph.runs[0]
    formula_s_A_run.bold = True

    doc.add_paragraph(f"Shared key calculated by B (s_B) is:", style='BodyText')
    formula_s_B = f"s_B = A^b mod p = {A}^{b} mod {p} = {s_B}"
    formula_s_B_paragraph = doc.add_paragraph(formula_s_B, style='Normal')
    formula_s_B_run = formula_s_B_paragraph.runs[0]
    formula_s_B_run.bold = True

    # 成功消息
    if s_A == s_B:
        success_paragraph = doc.add_paragraph(
            "Since s_A = s_B, the key exchange is successful and the shared key is the same.", style='BodyText')
        success_run = success_paragraph.runs[0]
        success_run.font.color.rgb = RGBColor(0, 153, 0)  # Green color
    else:
        fail_paragraph = doc.add_paragraph("The key exchange failed. The shared keys are different.", style='BodyText')
        fail_run = fail_paragraph.runs[0]
        fail_run.font.color.rgb = RGBColor(255, 0, 0)  # Red color

    # 添加额外美化
    doc.add_paragraph('\n' + '-' * 50 + '\n', style='Normal')
    doc.add_paragraph("This concludes the Diffie-Hellman key exchange process.", style='BodyText')

    # 添加页脚
    section = doc.sections[0]
    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer_paragraph.add_run("Diffie-Hellman Key Exchange - Confidential").italic = True

    # 保存到一个内存中的字节流
    byte_io = io.BytesIO()
    doc.save(byte_io)
    byte_io.seek(0)
    return byte_io

def main():
    st.set_page_config(page_title="Diffie-Hellman Key Exchange Demo", page_icon="🔑", layout="wide")

    # 添加背景色
    # st.markdown(
    #     """
    #     <style>
    #     .stApp {
    #         background-color: #2c3e50;
    #     }
    #     </style>
    #     """,
    #     unsafe_allow_html=True
    # )

    st.title('🔐 Diffie-Hellman Key Exchange Demo', anchor="top")
    # 添加创作者信息
    st.markdown("**小组成员**: 20221485 李昊轩， 20221587 史亚涛", unsafe_allow_html=True)
    st.markdown("**指导老师**: 陈欣", unsafe_allow_html=True)

    # Step 1: 选择素数p和原根g
    p = get_prime_in_range()
    g = find_primitive_root(p)

    # Step 2: 设置默认值或重新生成私钥
    if 'a_value' not in st.session_state or 'b_value' not in st.session_state:
        # 初始化随机数
        st.session_state.a_value = sympy.randprime(1, p - 1)
        st.session_state.b_value = sympy.randprime(1, p - 1)

    # Step 3: 用户输入私钥，并允许重新生成随机数
    a = st.number_input("Private key of A (a):", min_value=1, max_value=p - 1, value=st.session_state.a_value,
                        help="This is your private key A. Choose a random value between 1 and p-1.")
    b = st.number_input("Private key of B (b):", min_value=1, max_value=p - 1, value=st.session_state.b_value,
                        help="This is your private key B. Choose a random value between 1 and p-1.")

    # 显示 p 和 g
    st.markdown(f"<h4 style='color: #3498db;'>Prime p = {p}</h4>", unsafe_allow_html=True)
    st.markdown(f"<h4 style='color: #3498db;'>Primitive Root g = {g}</h4>", unsafe_allow_html=True)
    st.markdown(f"<h4 style='color: #3498db;'>Private key of A (a) = {a}</h4>", unsafe_allow_html=True)
    st.markdown(f"<h4 style='color: #3498db;'>Private key of B (b) = {b}</h4>", unsafe_allow_html=True)

    if st.button('🔄 Calculate Exchange'):
        # Step 4: 执行D-H密钥交换
        A, B, s_A, s_B = diffie_hellman_exchange(p, g, a, b)

        # 展示结果
        st.markdown(f"<h4 style='color: #9b59b6;'>Public key of A (A) = {A}</h4>", unsafe_allow_html=True)
        st.markdown(f"<h4 style='color: #9b59b6;'>Public key of B (B) = {B}</h4>", unsafe_allow_html=True)
        st.markdown(f"<h4 style='color: #9b59b6;'>Shared key calculated by A (s_A) = {s_A}</h4>",
                    unsafe_allow_html=True)
        st.markdown(f"<h4 style='color: #9b59b6;'>Shared key calculated by B (s_B) = {s_B}</h4>",
                    unsafe_allow_html=True)

        if s_A == s_B:
            st.success("✅ The key exchange is successful! The shared key is the same.")
        else:
            st.error("❌ The key exchange failed. The shared keys are different.")

        # Step 5: 生成Word文档
        byte_io = record_process_to_word(p, g, a, b, A, B, s_A, s_B)

        # 提供下载链接
        st.download_button(
            label="📥 Download Calculation Process as Word Document",
            data=byte_io,
            file_name="dh_key_exchange.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )


if __name__ == "__main__":
    main()
